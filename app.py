
import json
import os
import io
from flask import Flask, request, render_template_string, jsonify, send_file
from openai import OpenAI
from docx import Document

# --- 1. CONFIGURAÇÃO DA APLICAÇÃO FLASK E API ---
app = Flask(__name__)
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY")

# --- 2. TEMPLATE HTML COM GERADOR DOCX E INSTRUÇÕES KINDLE ---
HTML_TEMPLATE = '''
<!DOCTYPE html>
<html lang="pt-br">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Fábrica de Conhecimento CEO v7.0 (DOCX)</title>
    <style>
        body { font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif; background-color: #f0f2f5; color: #1c1e21; margin: 0; padding: 2rem; }
        .container { background: #fff; padding: 2rem 3rem; border-radius: 12px; box-shadow: 0 4px 12px rgba(0, 0, 0, 0.1); max-width: 800px; margin: auto; }
        h1, h3 { text-align: center; }
        h1 { font-size: 1.6rem; }
        h2 { font-size: 1.3rem; color: #1c1e21; margin-top: 1.5rem; border-bottom: 1px solid #dddfe2; padding-bottom: 0.5rem; }
        h3 { font-size: 1.1rem; color: #606770; margin-bottom: 2rem; }
        input[type="text"] { width: 100%; padding: 0.8rem; margin-bottom: 1rem; border: 1px solid #dddfe2; border-radius: 6px; font-size: 1rem; box-sizing: border-box; }
        
        .button-grid { display: grid; grid-template-columns: 1fr 1fr; gap: 0.5rem; margin-bottom: 1rem; }
        .block-button { background-color: #1877f2; color: white; border: none; padding: 0.8rem; border-radius: 6px; font-size: 0.9rem; cursor: pointer; transition: background-color 0.3s; }
        .block-button:hover:not(:disabled) { background-color: #166fe5; }
        .block-button:disabled { background-color: #dddfe2; color: #8a8d91; cursor: not-allowed; }
        .block-button.completed { background-color: #36a420; }

        .info { background-color: #e7f3ff; border-left: 5px solid #1877f2; padding: 1rem; margin: 1.5rem 0; text-align: left; font-size: 0.9rem; border-radius: 6px; }
        .footer { margin-top: 2rem; font-size: 0.8rem; color: #8a8d91; text-align: center; }
        #results { margin-top: 1rem; }
        .block-container { background-color: #fafafa; border: 1px solid #dddfe2; padding: 1.5rem; margin-top: 1rem; border-radius: 8px; }
        .block-content { white-space: pre-wrap; word-wrap: break-word; }

        #docx-button { background-color: #34495e; color: white; border: none; padding: 0.8rem 1.5rem; border-radius: 6px; font-size: 1rem; cursor: pointer; transition: background-color 0.3s; width: 100%; margin-top: 1.5rem; display: none; }
        #docx-button:hover:not(:disabled) { background-color: #2c3e50; }
        
        #kindle-instructions { background-color: #fdf9e4; border-left: 5px solid #f1c40f; padding: 1rem; margin-top: 1rem; font-size: 0.9rem; border-radius: 6px; display: none; }

        .placeholder { color: #8a8d91; font-style: italic; text-align: center; padding: 2rem; }
    </style>
</head>
<body>
    <div class="container">
        <h1>🏭 Fábrica de Conhecimento CEO</h1>
        <h3>Transformando livros em Dossiês Estratégicos</h3>
        
        <input type="text" id="livro-input" placeholder="1. Título do Livro (Obrigatório)" required>
        <input type="text" id="autor-input" placeholder="2. Autor (Opcional)">
        
        <div class="info"><strong>Estratégia:</strong> Gere cada bloco do dossiê individualmente.</div>
        
        <div class="button-grid">
            <button id="btn-block-1" class="block-button" onclick="generateBlock(1)">Gerar Bloco 1</button>
            <button id="btn-block-2" class="block-button" onclick="generateBlock(2)">Gerar Bloco 2</button>
            <button id="btn-block-3" class="block-button" onclick="generateBlock(3)">Gerar Bloco 3</button>
            <button id="btn-block-4" class="block-button" onclick="generateBlock(4)">Gerar Bloco 4</button>
        </div>

        <div id="results"><p class="placeholder">Os blocos gerados aparecerão aqui...</p></div>
        
        <button id="docx-button" onclick="generateDocx()">📥 Gerar Word (.docx)</button>
        <div id="kindle-instructions">
            <strong>Para enviar ao Kindle:</strong>
            <ol>
                <li>Baixe o arquivo <code>.docx</code> clicando no botão acima.</li>
                <li>Abra seu e-mail e anexe o arquivo que você baixou.</li>
                <li>Envie para seu endereço Kindle (ex: <code>seunome@kindle.com</code>).</li>
            </ol>
             <small>Você pode encontrar seu endereço Kindle e gerenciar e-mails aprovados na página "Gerencie seu conteúdo e dispositivos" na Amazon.</small>
        </div>

        <div class="footer">Desenvolvido para o Protocolo de Superpoder de Estudo CEO</div>
    </div>

    <script>
        async function generateBlock(blockNumber) {
            const livroInput = document.getElementById('livro-input');
            if (!livroInput.value) {
                alert("Por favor, insira o título do livro antes de gerar um bloco.");
                return;
            }

            const button = document.getElementById(`btn-block-${blockNumber}`);
            const resultsDiv = document.getElementById('results');

            button.disabled = true;
            button.textContent = "Gerando...";

            try {
                const response = await fetch('/generate-block', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({
                        livro: livroInput.value,
                        autor: document.getElementById('autor-input').value,
                        block_number: blockNumber
                    })
                });

                const data = await response.json();
                if (!response.ok) throw new Error(data.error || "Erro desconhecido no servidor.");
                
                const placeholder = resultsDiv.querySelector('.placeholder');
                if (placeholder) placeholder.remove();

                const blockContainer = document.createElement('div');
                blockContainer.className = 'block-container';
                blockContainer.id = `content-block-${blockNumber}`;

                blockContainer.innerHTML = `<h2>Parte ${data.indice}: ${data.tema}</h2><div class="block-content"></div>`;
                blockContainer.querySelector('.block-content').textContent = data.texto_bloco; // Use textContent for security

                const existingBlock = document.getElementById(blockContainer.id);
                if (existingBlock) resultsDiv.replaceChild(blockContainer, existingBlock);
                else resultsDiv.appendChild(blockContainer);

                button.textContent = `✅ Bloco ${blockNumber} Gerado`;
                button.classList.add('completed');
                document.getElementById('docx-button').style.display = 'block';
                document.getElementById('kindle-instructions').style.display = 'block';

            } catch (error) {
                alert(`Falha ao gerar o Bloco ${blockNumber}: ${error.message}`);
                button.disabled = false;
                button.textContent = `Gerar Bloco ${blockNumber}`;
            }
        }

        async function generateDocx() {
            const button = document.getElementById('docx-button');
            button.disabled = true;
            button.textContent = 'Gerando .docx...';

            const livro = document.getElementById('livro-input').value || "dossie";
            const autor = document.getElementById('autor-input').value;
            const blocks = [];
            document.querySelectorAll('.block-container').forEach(container => {
                blocks.push({
                    title: container.querySelector('h2').textContent,
                    content: container.querySelector('.block-content').textContent
                });
            });

            try {
                const response = await fetch('/generate-docx', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ livro, autor, blocks })
                });

                if (!response.ok) throw new Error((await response.json()).error || `Erro ${response.status}`);

                const blob = await response.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `DOSSIE_${livro.replace(/ /g, '_')}.docx`;
                document.body.appendChild(a);
                a.click();
                a.remove();
                window.URL.revokeObjectURL(url);

            } catch (error) {
                alert(`Falha ao gerar o arquivo: ${error.message}`);
            } finally {
                button.disabled = false;
                button.textContent = '📥 Gerar Word (.docx)';
            }
        }
    </script>
</body>
</html>
'''

# --- 3. LÓGICA DE GERAÇÃO PYTHON ---
def gerar_bloco_estrategico(client, nome_livro, autor_livro, tema, indice):
    prompt = f'''
    [ROLE] Atue como uma Cientista Experimental e CEO Estrategista.
    [OBJETIVO] Escreva a PARTE {indice} de um Dossiê Técnico exaustivo sobre o livro '{nome_livro}' de '{autor_livro}'.
    [FOCO DO BLOCO] {tema}.
    [DIRETRIZES TÉCNICAS]
    1. EXTENSÃO: Gere um texto com aproximadamente 3.000 palavras para este bloco.
    2. TOM: Executivo, analítico e focado em governança de alto nível.
    3. ESTRUTURA: Use parágrafos bem desenvolvidos. Não use markdown, apenas texto plano com quebras de linha.
    4. CONTEÚDO: Não resuma. Explique o 'Como' e o 'Porquê'. Conecte com EBITDA, cultura e escalabilidade.
    5. IDIOMA: Português, mantendo termos técnicos em Inglês entre parênteses.
    6. FOCO: Gere apenas o texto do bloco solicitado, sem introduções ou conclusões sobre o processo de geração.
    '''
    try:
        response = client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": "Você é um assistente que escreve conteúdo técnico detalhado em Português, usando apenas texto plano."},
                {"role": "user", "content": prompt}
            ],
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"ERRO_API: {str(e)}"

# --- 4. ROTAS DA APLICAÇÃO ---
@app.route('/')
def index():
    return render_template_string(HTML_TEMPLATE)

@app.route('/generate-block', methods=['POST'])
def generate_block_route():
    data = request.get_json()
    if not all([data.get('livro'), data.get('block_number')]):
        return jsonify({"error": "Dados insuficientes."}), 400
    if not DEEPSEEK_API_KEY: return jsonify({"error": "Chave da API não configurada."}), 500

    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url="https://api.deepseek.com/v1")
    blocos_temas = ["Engenharia de Sistemas e Tese Central Ontológica", "Mecanismos de Alavancagem e Casos de Estudo Reais", "Estratégias de Governança para Diretores e CEOs", "Protocolo de Implementação e Exercícios de Retenção (Cicatriz Sináptica)"]
    block_number = data['block_number']

    if not 1 <= block_number <= len(blocos_temas): return jsonify({"error": "Número do bloco inválido."}), 400

    tema = blocos_temas[block_number - 1]
    texto_gerado = gerar_bloco_estrategico(client, data['livro'], data.get('autor', ''), tema, block_number)

    if texto_gerado.startswith("ERRO_API:"): return jsonify({"error": texto_gerado}), 500

    return jsonify({"indice": block_number, "tema": tema, "texto_bloco": texto_gerado})

@app.route('/generate-docx', methods=['POST'])
def generate_docx_route():
    try:
        data = request.get_json()
        livro = data.get('livro', 'Dossiê')
        autor = data.get('autor', '')
        blocks = data.get('blocks', [])

        if not blocks: return jsonify({"error": "Nenhum conteúdo para gerar o documento."}), 400

        document = Document()
        document.add_heading(livro, level=0)
        if autor: document.add_paragraph(f'Por: {autor}')
        document.add_paragraph('')

        for block in blocks:
            document.add_heading(block.get('title', 'Seção'), level=1)
            document.add_paragraph(block.get('content', ''))

        f = io.BytesIO()
        document.save(f)
        f.seek(0)

        return send_file(f, as_attachment=True, download_name=f'DOSSIE_{livro.replace(" ", "_")}.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    except Exception as e:
        print(f"[SERVER ERROR] {e}")
        return jsonify({"error": "Erro interno no servidor ao gerar o arquivo DOCX."}), 500

if __name__ == "__main__":
    app.run(debug=True, threaded=True)
